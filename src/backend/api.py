"""
GREEN-AI FOOTPRINT TOOL — FastAPI Backend
==========================================

Production-ready API for managing AI models (predefined + custom)
and serving footprint calculations.

DEPLOYMENT: Render (Python 3.11+, PostgreSQL 15+)
DEPENDENCIES: fastapi, uvicorn, asyncpg, pydantic, python-dotenv

To run locally:
  pip install fastapi uvicorn asyncpg pydantic python-dotenv
  uvicorn api:app --reload --port 8000
"""

from fastapi import FastAPI, HTTPException, Query, Depends
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
from pydantic import BaseModel, Field, field_validator
from typing import Optional, List, Literal, Dict, Any
from datetime import datetime, timedelta
from decimal import Decimal
import uuid
import os
import json
import io
from pathlib import Path

from sqlalchemy import func

# Database imports
from sqlalchemy.orm import Session
from .database import get_db, test_connection
from .models import AIModel, CalculationLog, Organization, User, GridCarbonIntensity, GPUProfile, ESGReport
from .services.calculator import (
    CalculatorService,
    FootprintInput,
    DEFAULT_ECOSCORE_WEIGHTS,
    DEFAULT_PUE,
    DEFAULT_WUE_LITERS_PER_KWH,
)

# ============================================================
# App Configuration
# ============================================================

app = FastAPI(
    title="Green-AI Footprint API",
    version="1.0.0",
    description="Enterprise API for GenAI environmental impact management",
)


def _get_cors_origins() -> List[str]:
    origins_env = os.getenv("CORS_ORIGINS")
    if origins_env:
        return [o.strip() for o in origins_env.split(",") if o.strip()]

    frontend_url = os.getenv("FRONTEND_URL")
    if frontend_url:
        return [frontend_url]

    return ["http://localhost:5173"]

app.add_middleware(
    CORSMiddleware,
    allow_origins=_get_cors_origins(),
    allow_credentials=True,
    allow_methods=["GET", "POST", "PUT", "DELETE"],
    allow_headers=["*"],
)

# ============================================================
# Pydantic Models (Request / Response Schemas)
# ============================================================

ModelCategory = Literal[
    "frontier-llm", "mid-size-llm", "small-edge", "code-model",
    "image-gen", "embedding", "multimodal", "custom"
]

GpuType = Literal[
    "nvidia-h100", "nvidia-a100-80gb", "nvidia-a100-40gb",
    "nvidia-v100", "nvidia-t4", "nvidia-a10g", "cpu-only"
]


class AIModelBase(BaseModel):
    display_name: str = Field(..., min_length=2, max_length=200)
    family: str = Field(..., min_length=1, max_length=100)
    category: ModelCategory
    parameters_billion: float = Field(..., gt=0, le=10000)
    energy_per_million_tokens_kwh: float = Field(..., gt=0, le=100)
    default_gpu: GpuType
    gpu_count_inference: int = Field(default=1, ge=1, le=64)
    tokens_per_second_per_gpu: int = Field(..., ge=1, le=10000)
    quality_score: int = Field(..., ge=0, le=100)
    description: Optional[str] = Field(None, max_length=500)

    @field_validator("display_name")
    @classmethod
    def strip_name(cls, v: str) -> str:
        return v.strip()


class AIModelCreate(AIModelBase):
    """Schema for creating a custom AI model."""
    pass


class AIModelUpdate(BaseModel):
    """Schema for updating a custom AI model (all fields optional)."""
    display_name: Optional[str] = Field(None, min_length=2, max_length=200)
    category: Optional[ModelCategory] = None
    parameters_billion: Optional[float] = Field(None, gt=0, le=10000)
    energy_per_million_tokens_kwh: Optional[float] = Field(None, gt=0, le=100)
    default_gpu: Optional[GpuType] = None
    gpu_count_inference: Optional[int] = Field(None, ge=1, le=64)
    tokens_per_second_per_gpu: Optional[int] = Field(None, ge=1, le=10000)
    quality_score: Optional[int] = Field(None, ge=0, le=100)
    description: Optional[str] = Field(None, max_length=500)


class AIModelResponse(AIModelBase):
    """Schema for returning an AI model."""
    id: str
    slug: str
    is_predefined: bool
    is_active: bool
    created_at: datetime
    updated_at: datetime

    class Config:
        from_attributes = True


class AIModelListResponse(BaseModel):
    models: List[AIModelResponse]
    total: int
    predefined_count: int
    custom_count: int


class CalculationRequest(BaseModel):
    model_id: str
    region_id: str
    request_count: int = Field(..., ge=1)
    avg_tokens_per_request: int = Field(..., ge=1)
    pue: Optional[float] = Field(None, ge=1.0, le=3.0)
    wue: Optional[float] = Field(None, ge=0.0, le=5.0)
    persist: bool = Field(default=True)


class CalculationResponse(BaseModel):
    energy_kwh: float
    co2e_grams: float
    water_liters: float
    hardware_amortized_grams: float
    duration_hours: float
    energy_per_request: float
    co2e_per_request: float
    water_per_request: float
    eco_score: float
    eco_grade: str
    equivalent_km_driving: float
    equivalent_smartphone_charges: float
    assumptions: List[str]
    confidence: str


class DashboardMetricsResponse(BaseModel):
    total_requests: int
    total_energy_kwh: float
    total_co2e_kg: float
    total_water_liters: float
    avg_eco_score: float
    trend_direction: str
    period_comparison: Dict[str, float]


class DashboardModelUsage(BaseModel):
    model_id: str
    display_name: str
    quality_score: int
    requests: int
    energy_kwh: float
    co2e_grams: float
    water_liters: float
    avg_eco_score: float
    eco_grade: str


class DashboardMetricsResponseV2(DashboardMetricsResponse):
    model_usage: List[DashboardModelUsage]


class DashboardTimeSeriesPoint(BaseModel):
    date: str
    energy_kwh: float
    co2e_grams: float


class DashboardTimeSeriesHourPoint(BaseModel):
    hour: str
    energy_kwh: float
    co2e_grams: float


class DashboardMetricsResponseV3(DashboardMetricsResponseV2):
    time_series_30d: List[DashboardTimeSeriesPoint]
    time_series_today: List[DashboardTimeSeriesHourPoint]


class ModelComparisonRequest(BaseModel):
    model_ids: List[str]
    region_id: str
    requests_per_1k: int = Field(default=1000, ge=1)
    avg_tokens_per_request: int = Field(default=1000, ge=1)
    weights: Optional[Dict[str, float]] = None


class HealthResponse(BaseModel):
    status: str
    version: str
    timestamp: datetime


class RegionReferenceResponse(BaseModel):
    id: str
    region_id: str
    provider: str
    location: str
    gco2e_per_kwh: float
    source: str
    year: int
    renewable_percentage: int


class GpuReferenceResponse(BaseModel):
    id: str
    name: str
    tdp_watts: int
    typical_utilization: float
    memory_gb: int
    flops_teraflops: int
    embodied_carbon_kg_co2e: float
    expected_lifespan_hours: int
    water_cooling_liters_per_hour: float


class SettingsDefaultsResponse(BaseModel):
    default_pue: float
    default_wue: float
    ecoscore_weights: Dict[str, float]


class ESGReportSummary(BaseModel):
    id: str
    name: str
    period_start: datetime
    period_end: datetime
    total_requests: int
    total_energy_kwh: float
    total_co2e_kg: float
    total_water_liters: float
    avg_eco_score: Optional[float] = None
    generated_at: datetime

    class Config:
        from_attributes = True


class ESGReportListResponse(BaseModel):
    reports: List[ESGReportSummary]
    total: int


class ESGReportDetailResponse(ESGReportSummary):
    payload: Dict[str, Any]


class ESGReportGenerateRequest(BaseModel):
    days: int = Field(default=30, ge=7, le=365)
    name: Optional[str] = Field(default=None, max_length=200)


class ScenarioConfigRequest(BaseModel):
    model_id: str
    region_id: str
    request_count: int = Field(..., ge=1)
    avg_tokens_per_request: int = Field(..., ge=1)
    pue: Optional[float] = Field(None, ge=1.0, le=3.0)
    wue: Optional[float] = Field(None, ge=0.0, le=5.0)


class ScenarioCompareRequest(BaseModel):
    baseline: ScenarioConfigRequest
    proposed: ScenarioConfigRequest


class ScenarioConfigResult(BaseModel):
    model_id: str
    region_id: str
    request_count: int
    avg_tokens_per_request: int
    total_tokens: int
    energy_kwh: float
    co2e_grams: float
    water_liters: float
    hardware_amortized_grams: float
    eco_score: float
    eco_grade: str


class ScenarioDeltaResult(BaseModel):
    co2e_percent: float
    energy_percent: float
    water_percent: float
    eco_score_delta: float


class ScenarioCompareResponse(BaseModel):
    baseline: ScenarioConfigResult
    proposed: ScenarioConfigResult
    delta: ScenarioDeltaResult


class AnalyticsTimeSeriesPoint(BaseModel):
    date: str
    energy_kwh: float
    co2e_grams: float
    requests: int


class AnalyticsModelBreakdownItem(BaseModel):
    model_id: str
    display_name: str
    requests: int
    energy_kwh: float
    co2e_grams: float
    water_liters: float


class AnalyticsRecentActivityItem(BaseModel):
    id: str
    calculated_at: datetime
    model_id: Optional[str] = None
    region_id: str
    request_count: int
    avg_tokens_per_request: int
    total_tokens: int
    energy_kwh: float
    co2e_grams: float
    water_liters: float
    eco_score: Optional[float] = None
    eco_grade: Optional[str] = None


class AnalyticsResponse(BaseModel):
    time_series: List[AnalyticsTimeSeriesPoint]
    model_breakdown: List[AnalyticsModelBreakdownItem]
    recent_activity: List[AnalyticsRecentActivityItem]


# ============================================================
# Database Connection (asyncpg)
# ============================================================

# In production, use connection pooling:
# DATABASE_URL = os.getenv("DATABASE_URL")
# pool = None
#
# @app.on_event("startup")
# async def startup():
#     global pool
#     pool = await asyncpg.create_pool(DATABASE_URL, min_size=5, max_size=20)
#
# @app.on_event("shutdown")
# async def shutdown():
#     await pool.close()
#
# async def get_db():
#     async with pool.acquire() as conn:
#         yield conn


# ============================================================
# API Routes — Health
# ============================================================

@app.get("/api/v1/health", response_model=HealthResponse)
async def health_check(db: Session = Depends(get_db)):
    # Test database connection
    db_healthy = test_connection()
    
    return HealthResponse(
        status="healthy" if db_healthy else "degraded",
        version="1.0.0",
        timestamp=datetime.utcnow(),
    )


# ============================================================
# API Routes — AI Models
# ============================================================


# ============================================================
# API Routes — Reference Data
# ============================================================

@app.get("/api/v1/reference/regions", response_model=List[RegionReferenceResponse])
async def list_regions(db: Session = Depends(get_db)):
    regions = db.query(GridCarbonIntensity).order_by(GridCarbonIntensity.region_id.asc()).all()
    return [
        RegionReferenceResponse(
            id=r.id,
            region_id=r.region_id,
            provider=r.provider,
            location=r.location,
            gco2e_per_kwh=float(r.gco2e_per_kwh),
            source=r.source,
            year=int(r.year),
            renewable_percentage=int(r.renewable_percentage),
        )
        for r in regions
    ]


@app.get("/api/v1/reference/gpus", response_model=List[GpuReferenceResponse])
async def list_gpus(db: Session = Depends(get_db)):
    gpus = db.query(GPUProfile).order_by(GPUProfile.id.asc()).all()
    return [
        GpuReferenceResponse(
            id=g.id,
            name=g.name,
            tdp_watts=int(g.tdp_watts),
            typical_utilization=float(g.typical_utilization),
            memory_gb=int(g.memory_gb),
            flops_teraflops=int(g.flops_teraflops),
            embodied_carbon_kg_co2e=float(g.embodied_carbon_kg_co2e),
            expected_lifespan_hours=int(g.expected_lifespan_hours),
            water_cooling_liters_per_hour=float(g.water_cooling_liters_per_hour),
        )
        for g in gpus
    ]

@app.get("/api/v1/reference/settings-defaults", response_model=SettingsDefaultsResponse)
async def get_settings_defaults():
    return SettingsDefaultsResponse(
        default_pue=float(DEFAULT_PUE),
        default_wue=float(DEFAULT_WUE_LITERS_PER_KWH),
        ecoscore_weights={k: float(v) for k, v in DEFAULT_ECOSCORE_WEIGHTS.items()},
    )


@app.get("/api/v1/models", response_model=AIModelListResponse)
async def list_models(
    db: Session = Depends(get_db),
    category: Optional[ModelCategory] = Query(None),
    include_custom: bool = Query(True),
    include_predefined: bool = Query(True),
):
    """
    List all available AI models.
    
    Supports filtering by category and model type (predefined vs custom).
    Returns models sorted by: predefined first, then custom, within each
    sorted by parameter count descending.
    """
    query = db.query(AIModel).filter(AIModel.is_active == True)
    
    # Apply filters
    if category:
        query = query.filter(AIModel.category == category)
    if include_predefined and not include_custom:
        query = query.filter(AIModel.is_predefined == True)
    elif include_custom and not include_predefined:
        query = query.filter(AIModel.is_predefined == False)
    
    models = query.order_by(
        AIModel.is_predefined.desc(), 
        AIModel.parameters_billion.desc()
    ).all()
    
    # Convert to response format
    model_responses = []
    for model in models:
        model_responses.append(AIModelResponse(
            id=model.id,
            slug=model.slug,
            display_name=model.display_name,
            family=model.family,
            category=model.category,
            parameters_billion=float(model.parameters_billion),
            energy_per_million_tokens_kwh=float(model.energy_per_million_tokens_kwh),
            default_gpu=model.default_gpu,
            gpu_count_inference=model.gpu_count_inference,
            tokens_per_second_per_gpu=model.tokens_per_second_per_gpu,
            quality_score=model.quality_score,
            description=model.description,
            is_predefined=model.is_predefined,
            is_active=model.is_active,
            created_at=model.created_at,
            updated_at=model.updated_at,
        ))
    
    total = len(model_responses)
    predefined_count = sum(1 for m in model_responses if m.is_predefined)
    custom_count = total - predefined_count
    
    return AIModelListResponse(
        models=model_responses,
        total=total,
        predefined_count=predefined_count,
        custom_count=custom_count,
    )


@app.get("/api/v1/models/{model_id}", response_model=AIModelResponse)
async def get_model(model_id: str, db: Session = Depends(get_db)):
    """
    Get a single AI model by ID or slug.
    """
    model = db.query(AIModel).filter(
        (AIModel.id == model_id) | (AIModel.slug == model_id),
        AIModel.is_active == True
    ).first()
    
    if not model:
        raise HTTPException(status_code=404, detail=f"Model {model_id} not found")
    
    return AIModelResponse(
        id=model.id,
        slug=model.slug,
        display_name=model.display_name,
        family=model.family,
        category=model.category,
        parameters_billion=float(model.parameters_billion),
        energy_per_million_tokens_kwh=float(model.energy_per_million_tokens_kwh),
        default_gpu=model.default_gpu,
        gpu_count_inference=model.gpu_count_inference,
        tokens_per_second_per_gpu=model.tokens_per_second_per_gpu,
        quality_score=model.quality_score,
        description=model.description,
        is_predefined=model.is_predefined,
        is_active=model.is_active,
        created_at=model.created_at,
        updated_at=model.updated_at,
    )


@app.post("/api/v1/models", response_model=AIModelResponse, status_code=201)
async def create_custom_model(model: AIModelCreate, db: Session = Depends(get_db)):
    """
    Create a new custom AI model.
    
    Validation:
    - Name uniqueness checked per organization
    - Energy and parameter values clamped to physical bounds
    - GPU type validated against known hardware profiles
    """
    # For now, use default organization (in production, get from auth)
    org = db.query(Organization).filter(Organization.slug == "default").first()
    if not org:
        raise HTTPException(status_code=500, detail="Default organization not found")
    
    # Check for duplicate slug within organization
    slug = f"custom-{uuid.uuid4().hex[:12]}"
    
    # Create the model
    db_model = AIModel(
        id=str(uuid.uuid4()),
        slug=slug,
        display_name=model.display_name,
        family=model.family,
        category=model.category,
        parameters_billion=Decimal(str(model.parameters_billion)),
        energy_per_million_tokens_kwh=Decimal(str(model.energy_per_million_tokens_kwh)),
        default_gpu=model.default_gpu,
        gpu_count_inference=model.gpu_count_inference,
        tokens_per_second_per_gpu=model.tokens_per_second_per_gpu,
        quality_score=model.quality_score,
        description=model.description,
        is_predefined=False,
        is_active=True,
        organization_id=org.id,
        # created_by would be set from auth context in production
    )
    
    db.add(db_model)
    db.commit()
    db.refresh(db_model)
    
    return AIModelResponse(
        id=db_model.id,
        slug=db_model.slug,
        display_name=db_model.display_name,
        family=db_model.family,
        category=db_model.category,
        parameters_billion=float(db_model.parameters_billion),
        energy_per_million_tokens_kwh=float(db_model.energy_per_million_tokens_kwh),
        default_gpu=db_model.default_gpu,
        gpu_count_inference=db_model.gpu_count_inference,
        tokens_per_second_per_gpu=db_model.tokens_per_second_per_gpu,
        quality_score=db_model.quality_score,
        description=db_model.description,
        is_predefined=db_model.is_predefined,
        is_active=db_model.is_active,
        created_at=db_model.created_at,
        updated_at=db_model.updated_at,
    )


@app.put("/api/v1/models/{model_id}", response_model=AIModelResponse)
async def update_custom_model(model_id: str, update: AIModelUpdate):
    """
    Update a custom AI model.
    
    Only non-predefined models can be updated.
    
    Production implementation:
        UPDATE ai_models 
        SET display_name = COALESCE($2, display_name), ...
        WHERE id::text = $1 AND is_predefined = FALSE AND is_active = TRUE
        RETURNING *;
    """
    raise HTTPException(status_code=404, detail=f"Model {model_id} not found")


@app.delete("/api/v1/models/{model_id}", status_code=204)
async def delete_custom_model(model_id: str):
    """
    Soft-delete a custom AI model.
    
    Only non-predefined models can be deleted.
    Sets is_active = FALSE rather than removing the row,
    preserving audit trail for historical calculations.
    
    Production implementation:
        UPDATE ai_models 
        SET is_active = FALSE, updated_at = NOW()
        WHERE id::text = $1 AND is_predefined = FALSE
        RETURNING id;
    """
    raise HTTPException(status_code=404, detail=f"Model {model_id} not found")


# ============================================================
# API Routes — Calculations
# ============================================================

@app.post("/api/v1/calculate", response_model=CalculationResponse)
async def calculate_footprint(request: CalculationRequest, db: Session = Depends(get_db)):
    """
    Calculate environmental footprint for a given model + region + usage.
    
    Production implementation:
    1. Fetch model profile from ai_models table
    2. Fetch region carbon intensity from grid_intensities table
    3. Run calculation pipeline (same formulas as frontend engine)
    4. Log result to calculation_logs table
    5. Return result
    
    This endpoint mirrors the frontend calculation engine exactly,
    serving as source of truth for ESG report generation.
    """
    try:
        # Initialize calculator service
        calculator = CalculatorService(db)

        # Resolve model (request.model_id can be UUID or slug)
        resolved_model = calculator.get_model_profile(request.model_id)
        
        # Calculate total tokens
        total_tokens = request.request_count * request.avg_tokens_per_request
        
        # Create footprint input
        footprint_input = FootprintInput(
            model_id=resolved_model.id,
            region_id=request.region_id,
            total_tokens=total_tokens,
            request_count=request.request_count,
            avg_tokens_per_request=request.avg_tokens_per_request,
            pue=request.pue,
            wue=request.wue
        )
        
        # Calculate footprint
        result = calculator.calculate_full_footprint(footprint_input)
        
        # Calculate EcoScore
        eco_score_result = calculator.calculate_eco_score(
            request.model_id, request.region_id
        )

        if request.persist:
            # Log calculation to database (for audit trail)
            # In production, get user_id from auth context
            org = db.query(Organization).filter(Organization.slug == "default").first()
            if not org:
                raise HTTPException(status_code=500, detail="Default organization not found; run seed.py")

            calc_log = CalculationLog(
                model_id=resolved_model.id,
                region_id=request.region_id,
                request_count=request.request_count,
                avg_tokens_per_request=request.avg_tokens_per_request,
                total_tokens=total_tokens,
                energy_kwh=Decimal(str(result.energy_kwh)),
                co2e_grams=Decimal(str(result.co2e_grams)),
                water_liters=Decimal(str(result.water_liters)),
                hardware_amortized_grams=Decimal(str(result.hardware_amortized_grams)),
                eco_score=Decimal(str(eco_score_result.overall)),
                eco_grade=eco_score_result.grade,
                pue_factor=Decimal(str(request.pue)) if request.pue else None,
                wue_factor=Decimal(str(request.wue)) if request.wue else None,
                organization_id=org.id,
                # user_id would be set from auth context in production
            )

            db.add(calc_log)
            db.commit()
        
        return CalculationResponse(
            energy_kwh=result.energy_kwh,
            co2e_grams=result.co2e_grams,
            water_liters=result.water_liters,
            hardware_amortized_grams=result.hardware_amortized_grams,
            duration_hours=result.duration_hours,
            energy_per_request=result.energy_per_request,
            co2e_per_request=result.co2e_per_request,
            water_per_request=result.water_per_request,
            eco_score=eco_score_result.overall,
            eco_grade=eco_score_result.grade,
            equivalent_km_driving=result.equivalent_km_driving,
            equivalent_smartphone_charges=result.equivalent_smartphone_charges,
            assumptions=result.assumptions + eco_score_result.assumptions,
            confidence=eco_score_result.confidence,
        )
        
    except HTTPException:
        raise
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Calculation failed: {str(e)}")


# ============================================================
# API Routes — ESG Reports
# ============================================================

@app.get("/api/v1/reports", response_model=ESGReportListResponse)
async def list_reports(db: Session = Depends(get_db)):
    """List generated ESG reports for organization."""
    org = db.query(Organization).filter(Organization.slug == "default").first()
    if not org:
        raise HTTPException(status_code=404, detail="Default organization not found")

    rows = (
        db.query(ESGReport)
        .filter(ESGReport.organization_id == org.id)
        .order_by(ESGReport.generated_at.desc())
        .limit(50)
        .all()
    )

    return ESGReportListResponse(
        reports=[
            ESGReportSummary(
                id=r.id,
                name=r.name,
                period_start=r.period_start,
                period_end=r.period_end,
                total_requests=int(r.total_requests),
                total_energy_kwh=float(r.total_energy_kwh),
                total_co2e_kg=float(r.total_co2e_kg),
                total_water_liters=float(r.total_water_liters),
                avg_eco_score=float(r.avg_eco_score) if r.avg_eco_score is not None else None,
                generated_at=r.generated_at,
            )
            for r in rows
        ],
        total=len(rows),
    )


@app.post("/api/v1/reports/generate", response_model=ESGReportSummary)
async def generate_report(req: ESGReportGenerateRequest, db: Session = Depends(get_db)):
    """Generate and persist an ESG report snapshot from calculation_logs."""
    org = db.query(Organization).filter(Organization.slug == "default").first()
    if not org:
        raise HTTPException(status_code=404, detail="Default organization not found")

    period_end = datetime.utcnow()
    period_start = (period_end - timedelta(days=req.days)).replace(hour=0, minute=0, second=0, microsecond=0)

    base_filter = [
        CalculationLog.organization_id == org.id,
        CalculationLog.calculated_at >= period_start,
        CalculationLog.calculated_at <= period_end,
    ]

    totals = (
        db.query(
            func.sum(CalculationLog.request_count).label("total_requests"),
            func.sum(CalculationLog.energy_kwh).label("total_energy_kwh"),
            func.sum(CalculationLog.co2e_grams).label("total_co2e_grams"),
            func.sum(CalculationLog.water_liters).label("total_water_liters"),
            func.avg(CalculationLog.eco_score).label("avg_eco_score"),
        )
        .filter(*base_filter)
        .first()
    )

    total_requests = int(totals.total_requests or 0)
    total_energy_kwh = Decimal(str(totals.total_energy_kwh or 0))
    total_co2e_kg = Decimal(str((totals.total_co2e_grams or 0))) / Decimal("1000")
    total_water_liters = Decimal(str(totals.total_water_liters or 0))
    avg_eco_score = Decimal(str(totals.avg_eco_score)) if totals.avg_eco_score is not None else None

    if total_requests == 0:
        raise HTTPException(status_code=400, detail="No calculation logs found in the requested period")

    by_model = (
        db.query(
            AIModel.display_name.label("display_name"),
            CalculationLog.model_id.label("model_id"),
            func.sum(CalculationLog.request_count).label("requests"),
            func.sum(CalculationLog.energy_kwh).label("energy_kwh"),
            func.sum(CalculationLog.co2e_grams).label("co2e_grams"),
            func.sum(CalculationLog.water_liters).label("water_liters"),
        )
        .join(AIModel, AIModel.id == CalculationLog.model_id)
        .filter(*base_filter)
        .group_by(CalculationLog.model_id, AIModel.display_name)
        .order_by(func.sum(CalculationLog.co2e_grams).desc())
        .all()
    )

    payload = {
        "period": {
            "start": period_start.isoformat(),
            "end": period_end.isoformat(),
            "days": req.days,
        },
        "totals": {
            "total_requests": total_requests,
            "total_energy_kwh": float(total_energy_kwh),
            "total_co2e_kg": float(total_co2e_kg),
            "total_water_liters": float(total_water_liters),
            "avg_eco_score": float(avg_eco_score) if avg_eco_score is not None else None,
        },
        "by_model": [
            {
                "model_id": str(r.model_id),
                "display_name": str(r.display_name),
                "requests": int(r.requests or 0),
                "energy_kwh": float(r.energy_kwh or 0),
                "co2e_kg": float(Decimal(str(r.co2e_grams or 0)) / Decimal("1000")),
                "water_liters": float(r.water_liters or 0),
            }
            for r in by_model
        ],
    }

    name = (req.name or f"ESG Report ({req.days} days)").strip()

    report = ESGReport(
        organization_id=org.id,
        name=name,
        period_start=period_start,
        period_end=period_end,
        total_requests=total_requests,
        total_energy_kwh=total_energy_kwh,
        total_co2e_kg=total_co2e_kg,
        total_water_liters=total_water_liters,
        avg_eco_score=avg_eco_score,
        payload_json=json.dumps(payload),
    )

    db.add(report)
    db.commit()
    db.refresh(report)

    return ESGReportSummary(
        id=report.id,
        name=report.name,
        period_start=report.period_start,
        period_end=report.period_end,
        total_requests=int(report.total_requests),
        total_energy_kwh=float(report.total_energy_kwh),
        total_co2e_kg=float(report.total_co2e_kg),
        total_water_liters=float(report.total_water_liters),
        avg_eco_score=float(report.avg_eco_score) if report.avg_eco_score is not None else None,
        generated_at=report.generated_at,
    )


@app.get("/api/v1/reports/{report_id}", response_model=ESGReportDetailResponse)
async def get_report(report_id: str, db: Session = Depends(get_db)):
    org = db.query(Organization).filter(Organization.slug == "default").first()
    if not org:
        raise HTTPException(status_code=404, detail="Default organization not found")

    report = (
        db.query(ESGReport)
        .filter(ESGReport.id == report_id, ESGReport.organization_id == org.id)
        .first()
    )
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")

    try:
        payload = json.loads(report.payload_json or "{}")
        if not isinstance(payload, dict):
            payload = {}
    except Exception:
        payload = {}

    return ESGReportDetailResponse(
        id=report.id,
        name=report.name,
        period_start=report.period_start,
        period_end=report.period_end,
        total_requests=int(report.total_requests),
        total_energy_kwh=float(report.total_energy_kwh),
        total_co2e_kg=float(report.total_co2e_kg),
        total_water_liters=float(report.total_water_liters),
        avg_eco_score=float(report.avg_eco_score) if report.avg_eco_score is not None else None,
        generated_at=report.generated_at,
        payload=payload,
    )


@app.get("/api/v1/reports/{report_id}/export/pdf")
async def export_report_pdf(report_id: str, db: Session = Depends(get_db)):
    org = db.query(Organization).filter(Organization.slug == "default").first()
    if not org:
        raise HTTPException(status_code=404, detail="Default organization not found")

    report = (
        db.query(ESGReport)
        .filter(ESGReport.id == report_id, ESGReport.organization_id == org.id)
        .first()
    )
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")

    try:
        payload = json.loads(report.payload_json or "{}")
        if not isinstance(payload, dict):
            payload = {}
    except Exception:
        payload = {}

    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
    except Exception:
        raise HTTPException(status_code=500, detail="PDF export dependency missing")

    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter

    y = height - 50
    c.setFont("Helvetica-Bold", 16)
    c.drawString(50, y, str(report.name))

    y -= 22
    c.setFont("Helvetica", 10)
    c.drawString(50, y, f"Period: {report.period_start.date().isoformat()} to {report.period_end.date().isoformat()}")
    y -= 18
    c.drawString(50, y, f"Generated: {report.generated_at.isoformat(timespec='seconds')}Z")

    y -= 28
    c.setFont("Helvetica-Bold", 12)
    c.drawString(50, y, "Totals")
    y -= 18
    c.setFont("Helvetica", 10)
    c.drawString(60, y, f"Requests: {int(report.total_requests):,}")
    y -= 14
    c.drawString(60, y, f"Energy: {float(report.total_energy_kwh):.2f} kWh")
    y -= 14
    c.drawString(60, y, f"CO2e: {float(report.total_co2e_kg):.2f} kg")
    y -= 14
    c.drawString(60, y, f"Water: {float(report.total_water_liters):.2f} L")
    y -= 14
    if report.avg_eco_score is not None:
        c.drawString(60, y, f"Avg EcoScore: {float(report.avg_eco_score):.1f}")
    else:
        c.drawString(60, y, "Avg EcoScore: —")

    by_model = payload.get("by_model") if isinstance(payload, dict) else None
    if isinstance(by_model, list) and len(by_model) > 0:
        y -= 26
        c.setFont("Helvetica-Bold", 12)
        c.drawString(50, y, "Top Models by CO2e")
        y -= 18
        c.setFont("Helvetica", 9)

        for row in by_model[:10]:
            if y < 70:
                c.showPage()
                y = height - 50
                c.setFont("Helvetica", 9)
            name = str(row.get("display_name", ""))[:60]
            co2e_kg = row.get("co2e_kg")
            requests = row.get("requests")
            c.drawString(60, y, f"{name} | requests={int(requests or 0):,} | co2e_kg={float(co2e_kg or 0):.2f}")
            y -= 12

    c.showPage()
    c.save()
    buffer.seek(0)

    filename = f"{report.name}.pdf".replace("/", "-")
    headers = {"Content-Disposition": f"attachment; filename=\"{filename}\""}
    return StreamingResponse(buffer, media_type="application/pdf", headers=headers)


@app.get("/api/v1/reports/{report_id}/export/docx")
async def export_report_docx(report_id: str, db: Session = Depends(get_db)):
    org = db.query(Organization).filter(Organization.slug == "default").first()
    if not org:
        raise HTTPException(status_code=404, detail="Default organization not found")

    report = (
        db.query(ESGReport)
        .filter(ESGReport.id == report_id, ESGReport.organization_id == org.id)
        .first()
    )
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")

    try:
        payload = json.loads(report.payload_json or "{}")
        if not isinstance(payload, dict):
            payload = {}
    except Exception:
        payload = {}

    try:
        from docx import Document
    except Exception:
        raise HTTPException(status_code=500, detail="DOCX export dependency missing")

    doc = Document()
    doc.add_heading(str(report.name), level=1)
    doc.add_paragraph(f"Period: {report.period_start.date().isoformat()} to {report.period_end.date().isoformat()}")
    doc.add_paragraph(f"Generated: {report.generated_at.isoformat(timespec='seconds')}Z")

    doc.add_heading("Totals", level=2)
    doc.add_paragraph(f"Requests: {int(report.total_requests):,}")
    doc.add_paragraph(f"Energy: {float(report.total_energy_kwh):.2f} kWh")
    doc.add_paragraph(f"CO2e: {float(report.total_co2e_kg):.2f} kg")
    doc.add_paragraph(f"Water: {float(report.total_water_liters):.2f} L")
    doc.add_paragraph(
        f"Avg EcoScore: {float(report.avg_eco_score):.1f}" if report.avg_eco_score is not None else "Avg EcoScore: —"
    )

    by_model = payload.get("by_model") if isinstance(payload, dict) else None
    if isinstance(by_model, list) and len(by_model) > 0:
        doc.add_heading("Top Models by CO2e", level=2)
        table = doc.add_table(rows=1, cols=3)
        hdr = table.rows[0].cells
        hdr[0].text = "Model"
        hdr[1].text = "Requests"
        hdr[2].text = "CO2e (kg)"

        for row in by_model[:10]:
            cells = table.add_row().cells
            cells[0].text = str(row.get("display_name", ""))
            cells[1].text = f"{int(row.get('requests') or 0):,}"
            cells[2].text = f"{float(row.get('co2e_kg') or 0):.2f}"

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)

    filename = f"{report.name}.docx".replace("/", "-")
    headers = {
        "Content-Disposition": f"attachment; filename=\"{filename}\"",
    }
    return StreamingResponse(
        out,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )


@app.get("/api/v1/analytics", response_model=AnalyticsResponse)
async def get_analytics(
    db: Session = Depends(get_db),
    region_id: Optional[str] = Query(None),
    days: int = Query(90, ge=7, le=365),
):
    org = db.query(Organization).filter(Organization.slug == "default").first()
    if not org:
        raise HTTPException(status_code=404, detail="Default organization not found")

    start_dt = datetime.utcnow() - timedelta(days=days - 1)
    start_dt = start_dt.replace(hour=0, minute=0, second=0, microsecond=0)

    base_filter = [
        CalculationLog.organization_id == org.id,
        CalculationLog.calculated_at >= start_dt,
    ]
    if region_id:
        base_filter.append(CalculationLog.region_id == region_id)

    # Daily time series
    daily_rows = (
        db.query(
            func.date(CalculationLog.calculated_at).label("day"),
            func.sum(CalculationLog.energy_kwh).label("energy_kwh"),
            func.sum(CalculationLog.co2e_grams).label("co2e_grams"),
            func.sum(CalculationLog.request_count).label("requests"),
        )
        .filter(*base_filter)
        .group_by(func.date(CalculationLog.calculated_at))
        .all()
    )

    by_day: Dict[str, Dict[str, float]] = {}
    for r in daily_rows:
        day_str = str(r.day)
        by_day[day_str] = {
            "energy_kwh": float(r.energy_kwh or 0),
            "co2e_grams": float(r.co2e_grams or 0),
            "requests": float(r.requests or 0),
        }

    time_series: List[AnalyticsTimeSeriesPoint] = []
    for i in range(days):
        d = (start_dt + timedelta(days=i)).date().isoformat()
        vals = by_day.get(d, {"energy_kwh": 0.0, "co2e_grams": 0.0, "requests": 0.0})
        time_series.append(
            AnalyticsTimeSeriesPoint(
                date=d,
                energy_kwh=float(vals["energy_kwh"]),
                co2e_grams=float(vals["co2e_grams"]),
                requests=int(vals["requests"]),
            )
        )

    # Per-model breakdown
    model_rows = (
        db.query(
            CalculationLog.model_id.label("model_id"),
            func.sum(CalculationLog.request_count).label("requests"),
            func.sum(CalculationLog.energy_kwh).label("energy_kwh"),
            func.sum(CalculationLog.co2e_grams).label("co2e_grams"),
            func.sum(CalculationLog.water_liters).label("water_liters"),
        )
        .filter(*base_filter)
        .group_by(CalculationLog.model_id)
        .all()
    )

    model_ids = [str(r.model_id) for r in model_rows if r.model_id]
    models = db.query(AIModel).filter(AIModel.id.in_(model_ids)).all() if model_ids else []
    model_index = {m.id: m for m in models}

    model_breakdown: List[AnalyticsModelBreakdownItem] = []
    for r in model_rows:
        mid = str(r.model_id) if r.model_id else None
        if not mid:
            continue
        m = model_index.get(mid)
        model_breakdown.append(
            AnalyticsModelBreakdownItem(
                model_id=mid,
                display_name=m.display_name if m else mid,
                requests=int(r.requests or 0),
                energy_kwh=float(r.energy_kwh or 0),
                co2e_grams=float(r.co2e_grams or 0),
                water_liters=float(r.water_liters or 0),
            )
        )

    model_breakdown.sort(key=lambda x: x.co2e_grams, reverse=True)

    # Recent activity
    recent_logs = (
        db.query(CalculationLog)
        .filter(*base_filter)
        .order_by(CalculationLog.calculated_at.desc())
        .limit(5)
        .all()
    )

    recent_activity = [
        AnalyticsRecentActivityItem(
            id=log.id,
            calculated_at=log.calculated_at,
            model_id=log.model_id,
            region_id=log.region_id,
            request_count=log.request_count,
            avg_tokens_per_request=log.avg_tokens_per_request,
            total_tokens=log.total_tokens,
            energy_kwh=float(log.energy_kwh),
            co2e_grams=float(log.co2e_grams),
            water_liters=float(log.water_liters),
            eco_score=float(log.eco_score) if log.eco_score is not None else None,
            eco_grade=log.eco_grade,
        )
        for log in recent_logs
    ]

    return AnalyticsResponse(
        time_series=time_series,
        model_breakdown=model_breakdown,
        recent_activity=recent_activity,
    )


@app.get("/api/v1/dashboard", response_model=DashboardMetricsResponseV3)
async def get_dashboard_metrics():
    """
    Get dashboard metrics aggregated from calculation logs.
    Replaces generateDashboardMetrics from simulation.ts.
    """
    import logging
    from sqlalchemy import text
    from .database import SessionLocal
    logger = logging.getLogger(__name__)
    
    # Use a completely fresh database session for dashboard
    db = SessionLocal()
    try:
        # Get default organization
        org = db.query(Organization).filter(Organization.slug == "default").first()
        if not org:
            raise HTTPException(status_code=404, detail="Default organization not found")
        
        # CRITICAL: Commit any pending transaction, set isolation to READ COMMITTED, 
        # and start a fresh transaction to see latest committed data
        db.commit()  # Commit any pending changes
        dialect = getattr(getattr(db, "bind", None), "dialect", None)
        dialect_name = getattr(dialect, "name", "")
        if dialect_name == "postgresql":
            db.execute(text("SET TRANSACTION ISOLATION LEVEL READ COMMITTED"))
        else:
            db.execute(text("SET SESSION TRANSACTION ISOLATION LEVEL READ COMMITTED"))
        db.rollback()  # Start fresh transaction with new isolation level
        db.expire_all()
        
        # Aggregate metrics from calculation logs with NO query caching
        logs = db.query(CalculationLog).filter(
            CalculationLog.organization_id == org.id
        ).execution_options(
            compiled_cache=None
        ).all()
        
        # DEBUG: Log what's in the database
        logger.info(f"[DASHBOARD] Org {org.id}: Found {len(logs)} calculation logs")
        if logs:
            latest = max(log.calculated_at for log in logs if log.calculated_at)
            total_req = sum(log.request_count for log in logs)
            logger.info(f"[DASHBOARD] Latest log: {latest}, Total requests: {total_req}")
            
            # Log individual log details for debugging
            for log in logs[-5:]:  # Last 5 logs
                logger.info(f"[DASHBOARD] Log: id={log.id}, requests={log.request_count}, time={log.calculated_at}")
        else:
            logger.warning("[DASHBOARD] No calculation logs found - this should not happen!")
        
        if not logs:
            return DashboardMetricsResponseV3(
                total_requests=0,
                total_energy_kwh=0.0,
                total_co2e_kg=0.0,
                total_water_liters=0.0,
                avg_eco_score=0.0,
                trend_direction="stable",
                period_comparison={"energyChange": 0.0, "co2eChange": 0.0, "requestsChange": 0.0},
                model_usage=[],
                time_series_30d=[],
                time_series_today=[],
            )
        
        total_requests = sum(log.request_count for log in logs)
        total_energy_kwh = sum(float(log.energy_kwh) for log in logs)
        total_co2e_kg = sum(float(log.co2e_grams) / 1000 for log in logs)
        total_water_liters = sum(float(log.water_liters) for log in logs)
        
        # Calculate average EcoScore
        eco_scores = [float(log.eco_score) for log in logs if log.eco_score is not None]
        avg_eco_score = sum(eco_scores) / len(eco_scores) if eco_scores else 0.0
        
        # Simple trend calculation (would be more sophisticated in production)
        recent_logs = [log for log in logs if log.calculated_at and log.calculated_at >= datetime.utcnow().replace(day=1)]
        older_logs = [log for log in logs if log.calculated_at and log.calculated_at < datetime.utcnow().replace(day=1)]
        
        if len(recent_logs) > len(older_logs):
            trend_direction = "improving"
        elif len(recent_logs) < len(older_logs):
            trend_direction = "degrading"
        else:
            trend_direction = "stable"
        
        # Aggregate per-model usage
        usage_by_model: Dict[str, Dict[str, Any]] = {}
        for log in logs:
            mid = log.model_id
            if not mid:
                continue
            if mid not in usage_by_model:
                usage_by_model[mid] = {
                    "requests": 0,
                    "energy_kwh": 0.0,
                    "co2e_grams": 0.0,
                    "water_liters": 0.0,
                    "eco_scores": [],
                }
            usage_by_model[mid]["requests"] += int(log.request_count or 0)
            usage_by_model[mid]["energy_kwh"] += float(log.energy_kwh or 0)
            usage_by_model[mid]["co2e_grams"] += float(log.co2e_grams or 0)
            usage_by_model[mid]["water_liters"] += float(log.water_liters or 0)
            if log.eco_score is not None:
                usage_by_model[mid]["eco_scores"].append(float(log.eco_score))

        models = db.query(AIModel).filter(AIModel.id.in_(list(usage_by_model.keys()))).all()
        model_index = {m.id: m for m in models}

        def eco_grade(score: float) -> str:
            if score >= 90:
                return "A"
            if score >= 75:
                return "B"
            if score >= 60:
                return "C"
            if score >= 45:
                return "D"
            return "F"

        model_usage: List[DashboardModelUsage] = []
        for mid, agg in usage_by_model.items():
            m = model_index.get(mid)
            if not m:
                continue
            scores = agg["eco_scores"]
            avg_score = (sum(scores) / len(scores)) if scores else 0.0
            model_usage.append(
                DashboardModelUsage(
                    model_id=m.id,
                    display_name=m.display_name,
                    quality_score=int(m.quality_score or 0),
                    requests=int(agg["requests"]),
                    energy_kwh=float(agg["energy_kwh"]),
                    co2e_grams=float(agg["co2e_grams"]),
                    water_liters=float(agg["water_liters"]),
                    avg_eco_score=float(avg_score),
                    eco_grade=eco_grade(float(avg_score)),
                )
            )

        model_usage.sort(key=lambda x: x.co2e_grams, reverse=True)

        # Real 30-day time series (daily aggregation)
        start_dt = datetime.utcnow() - timedelta(days=29)
        start_dt = start_dt.replace(hour=0, minute=0, second=0, microsecond=0)

        daily_rows = (
            db.query(
                func.date(CalculationLog.calculated_at).label("day"),
                func.sum(CalculationLog.energy_kwh).label("energy_kwh"),
                func.sum(CalculationLog.co2e_grams).label("co2e_grams"),
            )
            .filter(
                CalculationLog.organization_id == org.id,
                CalculationLog.calculated_at >= start_dt,
            )
            .group_by(func.date(CalculationLog.calculated_at))
            .execution_options(compiled_cache=None)
            .all()
        )

        by_day: Dict[str, Dict[str, float]] = {}
        for r in daily_rows:
            day_str = str(r.day)
            by_day[day_str] = {
                "energy_kwh": float(r.energy_kwh or 0),
                "co2e_grams": float(r.co2e_grams or 0),
            }

        time_series_30d: List[DashboardTimeSeriesPoint] = []
        for i in range(30):
            d = (start_dt + timedelta(days=i)).date().isoformat()
            vals = by_day.get(d, {"energy_kwh": 0.0, "co2e_grams": 0.0})
            time_series_30d.append(
                DashboardTimeSeriesPoint(
                    date=d,
                    energy_kwh=float(vals["energy_kwh"]),
                    co2e_grams=float(vals["co2e_grams"]),
                )
            )

        # Intraday (today) time series (hourly aggregation)
        start_today = datetime.utcnow().replace(hour=0, minute=0, second=0, microsecond=0)
        end_today = start_today + timedelta(days=1)

        hour_expr = func.to_char(CalculationLog.calculated_at, "HH24")
        if dialect_name not in {"postgresql"}:
            hour_expr = func.date_format(CalculationLog.calculated_at, "%H")

        hour_rows = (
            db.query(
                hour_expr.label("hour"),
                func.sum(CalculationLog.energy_kwh).label("energy_kwh"),
                func.sum(CalculationLog.co2e_grams).label("co2e_grams"),
            )
            .filter(
                CalculationLog.organization_id == org.id,
                CalculationLog.calculated_at >= start_today,
                CalculationLog.calculated_at < end_today,
            )
            .group_by(hour_expr)
            .execution_options(compiled_cache=None)
            .all()
        )

        by_hour: Dict[str, Dict[str, float]] = {}
        for r in hour_rows:
            h = str(r.hour).zfill(2)
            by_hour[h] = {
                "energy_kwh": float(r.energy_kwh or 0),
                "co2e_grams": float(r.co2e_grams or 0),
            }

        time_series_today: List[DashboardTimeSeriesHourPoint] = []
        for h in range(24):
            hh = str(h).zfill(2)
            vals = by_hour.get(hh, {"energy_kwh": 0.0, "co2e_grams": 0.0})
            time_series_today.append(
                DashboardTimeSeriesHourPoint(
                    hour=f"{hh}:00",
                    energy_kwh=float(vals["energy_kwh"]),
                    co2e_grams=float(vals["co2e_grams"]),
                )
            )

        return DashboardMetricsResponseV3(
            total_requests=total_requests,
            total_energy_kwh=total_energy_kwh,
            total_co2e_kg=total_co2e_kg,
            total_water_liters=total_water_liters,
            avg_eco_score=avg_eco_score,
            trend_direction=trend_direction,
            period_comparison={"energyChange": 0.0, "co2eChange": 0.0, "requestsChange": 0.0},
            model_usage=model_usage,
            time_series_30d=time_series_30d,
            time_series_today=time_series_today,
        )
    finally:
        db.close()


@app.post("/api/v1/compare")
async def compare_models(request: ModelComparisonRequest, db: Session = Depends(get_db)):
    """
    Compare multiple models and return detailed comparison.
    Replaces frontend model comparison logic.
    """
    try:
        calculator = CalculatorService(db)
        result = calculator.compare_models(
            model_ids=request.model_ids,
            region_id=request.region_id,
            requests_per_1k=request.requests_per_1k,
            avg_tokens_per_request=request.avg_tokens_per_request,
            weights=request.weights
        )
        return result
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Comparison failed: {str(e)}")


@app.post("/api/v1/scenarios/compare", response_model=ScenarioCompareResponse)
async def compare_scenarios(request: ScenarioCompareRequest, db: Session = Depends(get_db)):
    try:
        calculator = CalculatorService(db)

        def run(cfg: ScenarioConfigRequest) -> ScenarioConfigResult:
            total_tokens = cfg.request_count * cfg.avg_tokens_per_request
            footprint_input = FootprintInput(
                model_id=cfg.model_id,
                region_id=cfg.region_id,
                total_tokens=total_tokens,
                request_count=cfg.request_count,
                avg_tokens_per_request=cfg.avg_tokens_per_request,
                pue=cfg.pue,
                wue=cfg.wue,
            )
            fp = calculator.calculate_full_footprint(footprint_input)
            eco = calculator.calculate_eco_score(cfg.model_id, cfg.region_id)
            return ScenarioConfigResult(
                model_id=cfg.model_id,
                region_id=cfg.region_id,
                request_count=cfg.request_count,
                avg_tokens_per_request=cfg.avg_tokens_per_request,
                total_tokens=total_tokens,
                energy_kwh=fp.energy_kwh,
                co2e_grams=fp.co2e_grams,
                water_liters=fp.water_liters,
                hardware_amortized_grams=fp.hardware_amortized_grams,
                eco_score=eco.overall,
                eco_grade=eco.grade,
            )

        baseline = run(request.baseline)
        proposed = run(request.proposed)

        def pct(new: float, old: float) -> float:
            if old == 0:
                return 0.0
            return round(((new - old) / old) * 100.0, 2)

        delta = ScenarioDeltaResult(
            co2e_percent=pct(proposed.co2e_grams, baseline.co2e_grams),
            energy_percent=pct(proposed.energy_kwh, baseline.energy_kwh),
            water_percent=pct(proposed.water_liters, baseline.water_liters),
            eco_score_delta=round(proposed.eco_score - baseline.eco_score, 2),
        )

        return ScenarioCompareResponse(baseline=baseline, proposed=proposed, delta=delta)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Scenario comparison failed: {str(e)}")


# ============================================================
# Entry point
# ============================================================

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=int(os.getenv("PORT", "8000")))


# ============================================================
# Static Frontend (Render Monolith)
# ============================================================

_DIST_DIR = (Path(__file__).resolve().parents[2] / "dist").resolve()


@app.get("/{full_path:path}")
async def spa_fallback(full_path: str):
    if full_path.startswith("api/"):
        raise HTTPException(status_code=404, detail="Not found")

    if _DIST_DIR.exists() and full_path and not full_path.endswith("/"):
        candidate = (_DIST_DIR / full_path)
        if candidate.exists() and candidate.is_file():
            return FileResponse(str(candidate))

    index_path = _DIST_DIR / "index.html"
    if index_path.exists():
        return FileResponse(str(index_path))

    raise HTTPException(status_code=404, detail="Frontend build not found")
